"""PDF processing utilities using PyMuPDF (fitz) and ReportLab."""

import base64
import io
from datetime import datetime, timezone
from typing import Any

import fitz  # PyMuPDF


def get_page_count(file_path: str) -> int:
    """Return the number of pages in a document file.

    Supports PDF via PyMuPDF. For DOCX/DOC files, falls back to 1
    (python-docx does not expose a reliable page count).
    """
    import os
    ext = os.path.splitext(file_path)[-1].lower()
    if ext in (".docx", ".doc"):
        if ext == ".docx":
            try:
                from docx import Document as DocxDocument
                docx_doc = DocxDocument(file_path)
                return max(1, len(docx_doc.sections))
            except Exception:
                return 1
        return 1  # .doc — no pure-Python page counter
    # Default: treat as PDF
    with fitz.open(file_path) as doc:
        return doc.page_count


def _load_font(size: int):
    """Load a system font at the given size, falling back to Pillow's default."""
    from PIL import ImageFont  # type: ignore

    font_paths = [
        # macOS
        "/System/Library/Fonts/Helvetica.ttc",
        "/System/Library/Fonts/Arial.ttf",
        "/Library/Fonts/Arial.ttf",
        # Linux (Debian/Ubuntu)
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
        # Linux (RHEL/Fedora)
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
    ]
    for path in font_paths:
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            continue
    return ImageFont.load_default()


def _render_docx_page(file_path: str, page_num: int = 0, dpi: int = 150) -> bytes:
    """Render a .docx document page as a PNG image using python-docx + Pillow.

    Extracts paragraph text from the document and lays it out on a white
    page canvas.  For multi-section documents the text is split approximately
    evenly across sections; page_num selects which section to render.
    """
    from PIL import Image, ImageDraw  # type: ignore

    width_px  = int(8.5 * dpi)
    height_px = int(11 * dpi)
    margin_px = int(1.0 * dpi)   # 1-inch margin
    line_spacing = int(dpi * 0.18)

    img  = Image.new("RGB", (width_px, height_px), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    body_font_size  = max(12, int(dpi * 0.12))
    title_font_size = max(16, int(dpi * 0.16))
    body_font  = _load_font(body_font_size)
    title_font = _load_font(title_font_size)

    # Extract paragraphs from the docx
    paragraphs: list[tuple[str, bool]] = []  # (text, is_heading)
    try:
        from docx import Document as DocxDocument  # type: ignore
        docx_doc = DocxDocument(file_path)
        for para in docx_doc.paragraphs:
            text = para.text.strip()
            if not text:
                paragraphs.append(("", False))
                continue
            is_heading = para.style.name.lower().startswith("heading")
            paragraphs.append((text, is_heading))
    except Exception:
        paragraphs = [("Document preview", True), ("(Content could not be extracted)", False)]

    if not paragraphs:
        paragraphs = [("(Empty document)", False)]

    # Split paragraphs into sections (one per page_num)
    # Simple strategy: divide paragraph list into equal chunks
    try:
        from docx import Document as DocxDocument  # type: ignore
        docx_doc = DocxDocument(file_path)
        num_sections = max(1, len(docx_doc.sections))
    except Exception:
        num_sections = 1

    chunk_size = max(1, len(paragraphs) // num_sections)
    start = page_num * chunk_size
    end   = start + chunk_size if page_num < num_sections - 1 else len(paragraphs)
    page_paragraphs = paragraphs[start:end]

    # Draw a light header bar with filename
    header_h = int(dpi * 0.35)
    draw.rectangle([(0, 0), (width_px, header_h)], fill=(240, 240, 248))
    import os as _os
    doc_name = _os.path.basename(file_path)
    header_font = _load_font(max(10, int(dpi * 0.09)))
    draw.text((margin_px, int(header_h * 0.25)), doc_name, fill=(100, 100, 140), font=header_font)

    # Render paragraphs
    y = header_h + line_spacing
    max_text_width = width_px - 2 * margin_px

    for text, is_heading in page_paragraphs:
        if y >= height_px - margin_px:
            break
        if not text:
            y += line_spacing
            continue

        font = title_font if is_heading else body_font
        color = (30, 30, 60) if is_heading else (50, 50, 50)

        # Word-wrap the text to fit within the page width
        words = text.split()
        line = ""
        for word in words:
            test_line = f"{line} {word}".strip()
            bbox = draw.textbbox((0, 0), test_line, font=font)
            if bbox[2] - bbox[0] > max_text_width and line:
                draw.text((margin_px, y), line, fill=color, font=font)
                line_bbox = draw.textbbox((0, 0), line, font=font)
                y += (line_bbox[3] - line_bbox[1]) + line_spacing
                line = word
                if y >= height_px - margin_px:
                    break
            else:
                line = test_line
        if line and y < height_px - margin_px:
            draw.text((margin_px, y), line, fill=color, font=font)
            line_bbox = draw.textbbox((0, 0), line, font=font)
            y += (line_bbox[3] - line_bbox[1]) + (line_spacing * 2 if is_heading else line_spacing)

    # Page number footer
    footer_font = _load_font(max(9, int(dpi * 0.08)))
    footer_text = f"Page {page_num + 1}"
    draw.text((margin_px, height_px - margin_px + int(dpi * 0.1)), footer_text, fill=(160, 160, 160), font=footer_font)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _render_doc_page(file_path: str, dpi: int = 150) -> bytes:
    """Render a legacy .doc file as a PNG using ReportLab to produce a PDF
    page in memory, then rasterize it with PyMuPDF.

    If the content cannot be extracted the function falls back to a styled
    placeholder that clearly identifies the document by filename.
    """
    import os as _os
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    doc_name = _os.path.basename(file_path)
    buf = io.BytesIO()
    pdf_doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Heading1"],
        fontSize=18,
        spaceAfter=12,
        textColor=colors.HexColor("#1a1a2e"),
    )
    body_style = ParagraphStyle(
        "DocBody",
        parent=styles["Normal"],
        fontSize=11,
        textColor=colors.HexColor("#444444"),
        spaceAfter=8,
    )
    note_style = ParagraphStyle(
        "DocNote",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.HexColor("#888888"),
        spaceAfter=4,
    )

    elements = [
        Paragraph(doc_name, title_style),
        Spacer(1, 0.2 * inch),
        Paragraph(
            "This document is in legacy .doc format. "
            "A full-fidelity preview requires Microsoft Word or LibreOffice.",
            body_style,
        ),
        Spacer(1, 0.1 * inch),
        Paragraph(
            "You can download the original file using the Download button above.",
            note_style,
        ),
    ]

    pdf_doc.build(elements)
    pdf_bytes = buf.getvalue()

    # Rasterize the first page of the generated PDF with PyMuPDF
    with fitz.open(stream=pdf_bytes, filetype="pdf") as tmp_doc:
        page = tmp_doc.load_page(0)
        zoom = dpi / 72
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        return pix.tobytes("png")


def _render_placeholder_page(dpi: int = 150) -> bytes:
    """Return a gray placeholder PNG as a last resort when no better renderer
    is available.  The image is roughly A4-proportioned.
    """
    width_px  = int(8.5 * dpi)
    height_px = int(11 * dpi)

    try:
        from PIL import Image, ImageDraw  # type: ignore

        img  = Image.new("RGB", (width_px, height_px), color=(240, 240, 240))
        draw = ImageDraw.Draw(img)

        msg_lines = ["Document preview", "not available"]
        font_size  = max(24, dpi // 6)
        font       = _load_font(font_size)

        line_heights = []
        line_widths  = []
        for line in msg_lines:
            bbox = draw.textbbox((0, 0), line, font=font)
            line_widths.append(bbox[2] - bbox[0])
            line_heights.append(bbox[3] - bbox[1])

        total_h = sum(line_heights) + (len(msg_lines) - 1) * int(font_size * 0.4)
        y = (height_px - total_h) // 2
        for i, line in enumerate(msg_lines):
            x = (width_px - line_widths[i]) // 2
            draw.text((x, y), line, fill=(120, 120, 120), font=font)
            y += line_heights[i] + int(font_size * 0.4)

        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()

    except ImportError:
        pass

    # Minimal hard-coded solid-gray PNG as absolute last resort.
    import struct
    import zlib

    def _png_chunk(tag: bytes, data: bytes) -> bytes:
        c = struct.pack(">I", len(data)) + tag + data
        return c + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)

    w, h = width_px, height_px
    ihdr     = struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0)
    raw_rows = b"".join(b"\x00" + bytes([0x99, 0x99, 0x99]) * w for _ in range(h))
    idat     = zlib.compress(raw_rows)
    return (
        b"\x89PNG\r\n\x1a\n"
        + _png_chunk(b"IHDR", ihdr)
        + _png_chunk(b"IDAT", idat)
        + _png_chunk(b"IEND", b"")
    )


def convert_doc_to_pdf(doc_path: str, output_pdf_path: str) -> bool:
    """Convert a .doc or .docx file to PDF and save it at output_pdf_path.

    For .docx files: uses python-docx to extract content and reportlab to build
    a PDF that accurately reflects the document's text content.

    For .doc files: tries LibreOffice subprocess first; falls back to a reportlab
    placeholder that clearly identifies the file and its legacy format.

    Args:
        doc_path: Absolute path to the source .doc/.docx file.
        output_pdf_path: Absolute path where the output PDF should be written.

    Returns:
        True if conversion succeeded, False if it failed completely.
    """
    import os as _os
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    ext = _os.path.splitext(doc_path)[-1].lower()
    doc_name = _os.path.basename(doc_path)

    try:
        if ext == ".docx":
            # Extract text from the docx and render it as a proper PDF via reportlab
            from docx import Document as DocxDocument  # type: ignore

            docx_doc = DocxDocument(doc_path)
            buf = io.BytesIO()
            pdf_doc = SimpleDocTemplate(
                buf,
                pagesize=letter,
                rightMargin=inch,
                leftMargin=inch,
                topMargin=inch,
                bottomMargin=inch,
            )
            styles = getSampleStyleSheet()
            heading_style = ParagraphStyle(
                "DocHeading",
                parent=styles["Heading1"],
                fontSize=14,
                spaceAfter=8,
                textColor=colors.HexColor("#1a1a2e"),
            )
            body_style = ParagraphStyle(
                "DocBody",
                parent=styles["Normal"],
                fontSize=10,
                spaceAfter=6,
                textColor=colors.HexColor("#222222"),
                leading=14,
            )

            elements: list = []
            for para in docx_doc.paragraphs:
                text = para.text.strip()
                if not text:
                    elements.append(Spacer(1, 0.08 * inch))
                    continue
                is_heading = para.style.name.lower().startswith("heading")
                # Escape reportlab special XML chars
                safe_text = (
                    text.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                )
                style = heading_style if is_heading else body_style
                elements.append(Paragraph(safe_text, style))

            if not elements:
                elements.append(Paragraph("(Empty document)", body_style))

            pdf_doc.build(elements)
            with open(output_pdf_path, "wb") as fh:
                fh.write(buf.getvalue())
            return True

        elif ext == ".doc":
            # Try LibreOffice for faithful .doc conversion
            import subprocess
            output_dir = _os.path.dirname(output_pdf_path)
            try:
                result = subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        output_dir,
                        doc_path,
                    ],
                    capture_output=True,
                    timeout=30,
                )
                if result.returncode == 0:
                    # LibreOffice names the output after the input filename stem
                    stem = _os.path.splitext(_os.path.basename(doc_path))[0]
                    lo_output = _os.path.join(output_dir, f"{stem}.pdf")
                    if _os.path.exists(lo_output):
                        if lo_output != output_pdf_path:
                            _os.rename(lo_output, output_pdf_path)
                        return True
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass  # LibreOffice not available; fall through to placeholder

            # Fallback: reportlab placeholder for legacy .doc
            buf = io.BytesIO()
            pdf_doc = SimpleDocTemplate(
                buf,
                pagesize=letter,
                rightMargin=inch,
                leftMargin=inch,
                topMargin=inch,
                bottomMargin=inch,
            )
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "DocTitle",
                parent=styles["Heading1"],
                fontSize=18,
                spaceAfter=12,
                textColor=colors.HexColor("#1a1a2e"),
            )
            body_style = ParagraphStyle(
                "DocBody",
                parent=styles["Normal"],
                fontSize=11,
                textColor=colors.HexColor("#444444"),
                spaceAfter=8,
            )
            note_style = ParagraphStyle(
                "DocNote",
                parent=styles["Normal"],
                fontSize=9,
                textColor=colors.HexColor("#888888"),
                spaceAfter=4,
            )
            elements = [
                Paragraph(doc_name, title_style),
                Spacer(1, 0.2 * inch),
                Paragraph(
                    "This document is in Legacy .doc format. "
                    "A full-fidelity preview requires Microsoft Word or LibreOffice.",
                    body_style,
                ),
                Spacer(1, 0.1 * inch),
                Paragraph(
                    "You can download the original file using the Download button.",
                    note_style,
                ),
            ]
            pdf_doc.build(elements)
            with open(output_pdf_path, "wb") as fh:
                fh.write(buf.getvalue())
            return True

    except Exception:
        return False

    return False


def render_page(file_path: str, page_num: int = 0, dpi: int = 150, preview_path: str | None = None) -> bytes:
    """Render a document page as PNG bytes.

    When preview_path is supplied (a pre-converted PDF for .doc/.docx files),
    it is used directly via PyMuPDF for full-fidelity rendering.  Otherwise the
    dispatch table below applies:

      - .pdf   → PyMuPDF (fitz) — full fidelity
      - .docx  → python-docx text extraction + Pillow layout
      - .doc   → ReportLab placeholder PDF rasterised by PyMuPDF
      - other  → generic Pillow placeholder

    Args:
        file_path: Path to the original document file.
        page_num: Zero-based page index.
        dpi: Rendering resolution (default 150).
        preview_path: Optional path to a pre-generated preview PDF.

    Returns:
        PNG image as bytes.
    """
    import os as _os

    # If a preview PDF exists, always use it for rendering (handles .doc/.docx perfectly)
    if preview_path and _os.path.exists(preview_path):
        with fitz.open(preview_path) as doc:
            if page_num < 0 or page_num >= doc.page_count:
                page_num = max(0, min(page_num, doc.page_count - 1))
            page = doc.load_page(page_num)
            zoom = dpi / 72
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            return pix.tobytes("png")

    ext = _os.path.splitext(file_path)[-1].lower()

    if ext == ".docx":
        try:
            return _render_docx_page(file_path, page_num, dpi)
        except Exception:
            return _render_placeholder_page(dpi)

    if ext == ".doc":
        try:
            return _render_doc_page(file_path, dpi)
        except Exception:
            return _render_placeholder_page(dpi)

    # Default: treat as PDF via PyMuPDF
    with fitz.open(file_path) as doc:
        if page_num < 0 or page_num >= doc.page_count:
            raise ValueError(
                f"Page {page_num} is out of range for document with {doc.page_count} pages"
            )
        page = doc.load_page(page_num)
        zoom = dpi / 72
        mat  = fitz.Matrix(zoom, zoom)
        pix  = page.get_pixmap(matrix=mat, alpha=False)
        return pix.tobytes("png")


def apply_fields_to_pdf(file_path: str, fields: list[dict[str, Any]]) -> bytes:
    """Overlay field values onto the PDF and return the modified PDF bytes.

    Args:
        file_path: Path to the source PDF.
        fields: List of field dicts with keys: type, page, x, y, width, height, value, label.

    Returns:
        Modified PDF as bytes.
    """
    with fitz.open(file_path) as doc:
        for field in fields:
            if not field.get("value"):
                continue

            page_num = (field.get("page") or 1) - 1
            if page_num < 0 or page_num >= doc.page_count:
                continue

            page = doc.load_page(page_num)
            page_rect = page.rect

            # Field coordinates are stored as percentages of page dimensions.
            # Convert from % to PDF points using the actual page rectangle.
            x_pct = float(field.get("x", 0))
            y_pct = float(field.get("y", 0))
            w_pct = float(field.get("width", 10))
            h_pct = float(field.get("height", 5))

            pw = page_rect.width
            ph = page_rect.height
            x = pw * x_pct / 100.0
            y = ph * y_pct / 100.0
            w = pw * w_pct / 100.0
            h = ph * h_pct / 100.0

            rect = fitz.Rect(x, y, x + w, y + h)
            value = str(field["value"])
            field_type = field.get("type", "text")

            if field_type in ("signature", "initial"):
                if value.startswith("data:image"):
                    try:
                        header, encoded = value.split(",", 1)
                        img_bytes = base64.b64decode(encoded)
                        page.insert_image(rect, stream=img_bytes, keep_proportion=True)
                    except Exception:
                        page.insert_textbox(rect, "[ Signature ]", fontsize=14, color=(0.5, 0.5, 0.5), align=1)
                else:
                    # Fallback: draw as styled text
                    page.insert_textbox(
                        rect,
                        value,
                        fontsize=18,
                        color=(0.1, 0.2, 0.6),
                        align=1,
                    )
            elif field_type == "date_signed":
                page.insert_textbox(
                    rect,
                    value,
                    fontsize=10,
                    color=(0, 0, 0),
                    align=0,
                )
            elif field_type == "checkbox":
                if value.lower() in ("true", "1", "yes", "checked"):
                    # Draw checkmark
                    page.draw_rect(rect, color=(0, 0, 0), width=1)
                    p1 = fitz.Point(x + w * 0.2, y + h * 0.5)
                    p2 = fitz.Point(x + w * 0.45, y + h * 0.75)
                    p3 = fitz.Point(x + w * 0.8, y + h * 0.25)
                    page.draw_polyline([p1, p2, p3], color=(0, 0.6, 0), width=2)
            else:
                page.insert_textbox(
                    rect,
                    value,
                    fontsize=10,
                    color=(0, 0, 0),
                    align=0,
                )

        return doc.tobytes(garbage=4, deflate=True)


def generate_certificate(envelope_data: dict[str, Any], audit_events: list[dict[str, Any]]) -> bytes:
    """Generate a completion certificate PDF for a signed envelope.

    Args:
        envelope_data: Dict with envelope metadata (subject, id, sent_at, completed_at).
        audit_events: List of audit event dicts.

    Returns:
        PDF certificate as bytes.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        HRFlowable,
    )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CertTitle",
        parent=styles["Heading1"],
        fontSize=24,
        spaceAfter=12,
        textColor=colors.HexColor("#1a1a2e"),
    )
    subtitle_style = ParagraphStyle(
        "CertSubtitle",
        parent=styles["Normal"],
        fontSize=12,
        textColor=colors.HexColor("#4a4a6a"),
        spaceAfter=6,
    )
    label_style = ParagraphStyle(
        "Label",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.HexColor("#666666"),
        spaceAfter=2,
    )
    value_style = ParagraphStyle(
        "Value",
        parent=styles["Normal"],
        fontSize=11,
        textColor=colors.HexColor("#1a1a2e"),
        spaceAfter=8,
    )

    elements = []

    elements.append(Paragraph("Certificate of Completion", title_style))
    elements.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#1a73e8")))
    elements.append(Spacer(1, 0.2 * inch))

    elements.append(Paragraph("Document Details", styles["Heading2"]))
    elements.append(Spacer(1, 0.1 * inch))

    subject = envelope_data.get("subject", "N/A")
    env_id = str(envelope_data.get("id", "N/A"))
    completed_at = envelope_data.get("completed_at")
    sent_at = envelope_data.get("sent_at")

    details_data = [
        ["Subject:", subject],
        ["Envelope ID:", env_id],
        ["Sent:", str(sent_at) if sent_at else "N/A"],
        ["Completed:", str(completed_at) if completed_at else "N/A"],
        ["Certificate Generated:", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")],
    ]
    details_table = Table(details_data, colWidths=[2 * inch, 4.5 * inch])
    details_table.setStyle(
        TableStyle(
            [
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#666666")),
                ("TEXTCOLOR", (1, 0), (1, -1), colors.HexColor("#1a1a2e")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    elements.append(details_table)
    elements.append(Spacer(1, 0.3 * inch))

    if audit_events:
        elements.append(Paragraph("Audit Trail", styles["Heading2"]))
        elements.append(Spacer(1, 0.1 * inch))

        audit_data = [["Timestamp", "Event", "Details"]]
        for event in audit_events:
            ts = str(event.get("created_at", ""))[:19]
            evt = str(event.get("event_type", ""))
            details = str(event.get("details") or "")[:60]
            audit_data.append([ts, evt, details])

        audit_table = Table(audit_data, colWidths=[2 * inch, 2 * inch, 2.5 * inch])
        audit_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a73e8")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8f9fa")]),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#dadce0")),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        elements.append(audit_table)

    doc.build(elements)
    return buffer.getvalue()
