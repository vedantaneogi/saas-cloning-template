import csv
import io
import logging
import os
import shutil
import uuid
from typing import List, Optional

logger = logging.getLogger(__name__)

from fastapi import (
    APIRouter,
    Depends,
    File,
    Form,
    HTTPException,
    Query,
    UploadFile,
    status,
)
from fastapi.responses import Response as FastAPIResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.auth.dependencies import get_current_user
from app.auth.models import User
from app.core.config import get_settings
from app.db.session import get_db
from app.envelopes import service as svc
from app.envelopes.models import (
    Document,
    Envelope,
    EnvelopeStatus,
)
from app.envelopes.schemas import (
    BulkSendResponse,
    BulkSendStatusResponse,
    CommentCreate,
    CommentResponse,
    DocumentResponse,
    EnvelopeCreate,
    EnvelopeDetailResponse,
    EnvelopeListItem,
    EnvelopeListResponse,
    EnvelopeResponse,
    EnvelopeUpdate,
    FieldBulkSaveRequest,
    FieldCreate,
    FieldResponse,
    FieldUpdate,
    FolderCreate,
    FolderResponse,
    MoveEnvelopesRequest,
    RecipientCreate,
    RecipientResponse,
    RecipientUpdate,
    SaveAsTemplateRequest,
    SaveAsTemplateResponse,
    VoidEnvelopeRequest,
)
from app.signing.pdf_processor import convert_doc_to_pdf, get_page_count, render_page

settings = get_settings()

router = APIRouter(tags=["envelopes"])


# ── Envelope endpoints ────────────────────────────────────────────────────────

@router.post("/envelopes", response_model=EnvelopeResponse, status_code=status.HTTP_201_CREATED)
async def create_envelope(
    data: EnvelopeCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    envelope = await svc.create_envelope(db, current_user.id, data)
    return envelope


@router.get("/envelopes", response_model=EnvelopeListResponse)
async def list_envelopes(
    status_filter: Optional[EnvelopeStatus] = Query(None, alias="status"),
    # Virtual filter: "waiting_for_others" means envelopes the sender is waiting on
    # (status = sent OR delivered, not yet completed).
    envelope_filter: Optional[str] = Query(None, alias="filter"),
    search: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    date_from: Optional[str] = Query(None),
    date_to: Optional[str] = Query(None),
    envelope_id: Optional[str] = Query(None),
    recipient_search: Optional[str] = Query(None),
    shared: Optional[bool] = Query(None),
    folder_id: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    statuses: Optional[list[EnvelopeStatus]] = None
    moved_to_filter: Optional[str] = None
    expiring_soon: bool = False
    action_required: bool = False
    auth_failed: bool = False
    inbox: bool = False

    if envelope_filter == "inbox":
        inbox = True
        statuses = [EnvelopeStatus.sent, EnvelopeStatus.delivered, EnvelopeStatus.completed, EnvelopeStatus.declined, EnvelopeStatus.voided]
        status_filter = None
    elif envelope_filter == "waiting_for_others":
        # "Waiting for Others" — envelopes the current user sent that recipients
        # have not yet fully signed.  Covers both "sent" (awaiting first open) and
        # "delivered" (opened but not yet completed).
        statuses = [EnvelopeStatus.sent, EnvelopeStatus.delivered]
        status_filter = None  # ignore any concurrent status param
    elif envelope_filter == "deleted":
        # "Deleted" — soft-deleted envelopes (moved_to = "deleted")
        moved_to_filter = "deleted"
        status_filter = None
    elif envelope_filter == "expiring_soon":
        # "Expiring Soon" — envelopes with expires_at within 30 days
        expiring_soon = True
        status_filter = None
    elif envelope_filter == "action_required":
        # "Action Required" — envelopes where the current user is a recipient
        # who needs to take action (sign/approve)
        action_required = True
        status_filter = None
    elif envelope_filter == "auth_failed":
        # "Authentication Failed" — envelopes with recipients who have access
        # codes set (potential auth gate). This is a best-effort filter since
        # we don't currently track failed auth attempts per-envelope.
        auth_failed = True
        status_filter = None
    elif envelope_filter == "sent":
        # "Sent" folder — all envelopes authored by current user that are in
        # progress or have been declined/voided.
        statuses = [EnvelopeStatus.sent, EnvelopeStatus.delivered, EnvelopeStatus.declined, EnvelopeStatus.voided]
        status_filter = None

    items, total = await svc.list_envelopes(
        db,
        current_user.id,
        status=status_filter,
        statuses=statuses,
        search=search,
        page=page,
        page_size=page_size,
        date_from=date_from,
        date_to=date_to,
        envelope_id=envelope_id,
        recipient_search=recipient_search,
        shared=shared,
        current_user_email=current_user.email,
        folder_id=folder_id,
        moved_to=moved_to_filter,
        expiring_soon=expiring_soon,
        action_required=action_required,
        auth_failed=auth_failed,
        inbox=inbox,
    )
    pages = max(1, -(-total // page_size))  # ceiling division

    # Build list items, injecting from_name / from_email from the eager-loaded owner
    serialized = []
    for env in items:
        item = EnvelopeListItem.model_validate(env)
        if hasattr(env, "owner") and env.owner is not None:
            item.from_name = env.owner.name
            item.from_email = env.owner.email
        serialized.append(item)

    return EnvelopeListResponse(items=serialized, total=total, page=page, page_size=page_size, pages=pages)


@router.get("/envelopes/stats")
async def get_envelope_stats(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    from datetime import datetime, timezone as _tz, timedelta
    from sqlalchemy import func, select, or_
    from app.envelopes.models import Recipient

    result = await db.execute(
        select(Envelope.status, func.count(Envelope.id))
        .where(
            Envelope.user_id == current_user.id,
            or_(Envelope.moved_to.is_(None), Envelope.moved_to != "deleted"),
        )
        .group_by(Envelope.status)
    )
    stats = {row[0]: row[1] for row in result.all()}
    sent_count = stats.get(EnvelopeStatus.sent, 0) + stats.get(EnvelopeStatus.delivered, 0)

    waiting_result = await db.execute(
        select(func.count(func.distinct(Envelope.id)))
        .where(
            Envelope.user_id == current_user.id,
            Envelope.status.in_([EnvelopeStatus.sent, EnvelopeStatus.delivered]),
            or_(Envelope.moved_to.is_(None), Envelope.moved_to != "deleted"),
        )
    )
    waiting_for_others = waiting_result.scalar() or 0

    now = datetime.now(_tz.utc)
    thirty_days = now + timedelta(days=30)
    expiring_result = await db.execute(
        select(func.count(func.distinct(Envelope.id)))
        .where(
            Envelope.user_id == current_user.id,
            Envelope.expires_at.isnot(None),
            Envelope.expires_at > now,
            Envelope.expires_at <= thirty_days,
            Envelope.status.in_([EnvelopeStatus.sent, EnvelopeStatus.delivered]),
            or_(Envelope.moved_to.is_(None), Envelope.moved_to != "deleted"),
        )
    )
    expiring_soon = expiring_result.scalar() or 0

    return {
        "total": sum(stats.values()),
        "draft": stats.get(EnvelopeStatus.draft, 0),
        "sent": sent_count,
        "completed": stats.get(EnvelopeStatus.completed, 0),
        "voided": stats.get(EnvelopeStatus.voided, 0),
        "declined": stats.get(EnvelopeStatus.declined, 0),
        "waitingForOthers": waiting_for_others,
        "expiringSoon": expiring_soon,
    }


@router.get("/envelopes/tasks")
async def get_envelope_tasks(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Return envelopes where the current user is a pending recipient (needs to act).

    Includes the sender's display name and recipient progress counts so the UI
    can render a proper task card and signing progress bar.
    """
    from sqlalchemy import select, func, or_
    from sqlalchemy.orm import selectinload as _selectinload
    from app.envelopes.models import Recipient, RecipientRole, RecipientStatus

    # Step 1: find envelope IDs where the current user is an unsigned signer.
    # After send(), the first routing group gets status "sent"; later groups stay
    # "pending". We include all three pre-signed statuses so we don't miss anyone.
    action_id_result = await db.execute(
        select(Envelope.id, Envelope.updated_at)
        .join(Recipient, Recipient.envelope_id == Envelope.id)
        .where(
            func.lower(Recipient.email) == current_user.email.lower(),
            Recipient.role == RecipientRole.signer,
            Recipient.status.in_([RecipientStatus.pending, RecipientStatus.sent, RecipientStatus.delivered]),
            Envelope.status.in_([EnvelopeStatus.sent, EnvelopeStatus.delivered]),
            or_(Envelope.moved_to.is_(None), Envelope.moved_to != "deleted"),
        )
        .order_by(Envelope.updated_at.desc())
        .limit(10)
        .distinct()
    )
    action_ids = [row[0] for row in action_id_result.all()]

    # Step 2: load those envelopes with recipients and owner eagerly
    action_envs: list[Envelope] = []
    if action_ids:
        env_result = await db.execute(
            select(Envelope)
            .where(Envelope.id.in_(action_ids))
            .options(
                _selectinload(Envelope.recipients),
                _selectinload(Envelope.owner),
            )
            .order_by(Envelope.updated_at.desc())
        )
        action_envs = env_result.scalars().all()

    tasks = []
    for envelope in action_envs:
        signers = [r for r in envelope.recipients if r.role == RecipientRole.signer]
        completed = [r for r in signers if r.status == RecipientStatus.signed]
        sender_name = (
            envelope.owner.name
            if hasattr(envelope, "owner") and envelope.owner
            else current_user.name
        )
        tasks.append({
            "id": str(envelope.id),
            "subject": envelope.subject,
            "status": envelope.status,
            "from_name": sender_name,
            "action": "Needs to Sign",
            "recipients_total": len(signers),
            "recipients_completed": len(completed),
            "created_at": envelope.created_at.isoformat() if envelope.created_at else None,
            "updated_at": envelope.updated_at.isoformat() if envelope.updated_at else None,
        })

    # Also include envelopes the user sent that are awaiting others (secondary tasks)
    remaining = max(0, 10 - len(tasks))
    if remaining > 0:
        sent_result = await db.execute(
            select(Envelope)
            .options(
                _selectinload(Envelope.recipients),
            )
            .where(
                Envelope.user_id == current_user.id,
                Envelope.status.in_([EnvelopeStatus.sent, EnvelopeStatus.delivered]),
                or_(Envelope.moved_to.is_(None), Envelope.moved_to != "deleted"),
            )
            .order_by(Envelope.updated_at.desc())
            .limit(remaining)
        )
        sent_envelopes = sent_result.scalars().all()

        # Avoid duplicates (envelope might appear in both queries if user sent to themselves)
        existing_ids = {t["id"] for t in tasks}
        for envelope in sent_envelopes:
            if str(envelope.id) in existing_ids:
                continue
            signers = [r for r in envelope.recipients if r.role == RecipientRole.signer]
            completed = [r for r in signers if r.status == RecipientStatus.signed]
            tasks.append({
                "id": str(envelope.id),
                "subject": envelope.subject,
                "status": envelope.status,
                "from_name": current_user.name,
                "action": "Waiting for Others",
                "recipients_total": len(signers),
                "recipients_completed": len(completed),
                "created_at": envelope.created_at.isoformat() if envelope.created_at else None,
                "updated_at": envelope.updated_at.isoformat() if envelope.updated_at else None,
            })

    return tasks


# NOTE: Static sub-paths (bulk-send, move, etc.) MUST be registered before the
# wildcard /{envelope_id} route so FastAPI doesn't swallow them.

@router.put("/envelopes/move", status_code=status.HTTP_200_OK)
async def move_envelopes(
    data: MoveEnvelopesRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Move one or more envelopes to a folder or virtual view (inbox/sent)."""
    count = await svc.move_envelopes(db, current_user.id, data.envelope_ids, data.folder_id, data.moved_to)
    return {"moved": count}


@router.get("/envelopes/bulk-send/{batch_id}/status", response_model=BulkSendStatusResponse)
async def get_bulk_send_status_early(
    batch_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Return progress counts for a bulk send batch."""
    from sqlalchemy import func as sa_func, select as sa_select

    result = await db.execute(
        sa_select(Envelope.status, sa_func.count(Envelope.id))
        .where(Envelope.user_id == current_user.id, Envelope.batch_id == batch_id)
        .group_by(Envelope.status)
    )
    counts: dict[str, int] = {row[0]: row[1] for row in result.all()}
    total = sum(counts.values())
    if total == 0:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Batch not found")

    return BulkSendStatusResponse(
        batch_id=batch_id,
        total=total,
        sent=counts.get(EnvelopeStatus.sent, 0) + counts.get(EnvelopeStatus.delivered, 0),
        completed=counts.get(EnvelopeStatus.completed, 0),
        failed=counts.get(EnvelopeStatus.voided, 0) + counts.get(EnvelopeStatus.declined, 0),
    )


@router.get("/envelopes/bulk-send/batches")
async def list_bulk_batches(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    from sqlalchemy import func, select
    result = await db.execute(
        select(
            Envelope.batch_id,
            func.min(Envelope.subject).label("name"),
            func.count(Envelope.id).label("total"),
            func.count(Envelope.id).filter(Envelope.status.in_([EnvelopeStatus.sent, EnvelopeStatus.delivered])).label("sent"),
            func.count(Envelope.id).filter(Envelope.status == EnvelopeStatus.completed).label("completed"),
            func.count(Envelope.id).filter(Envelope.status.in_([EnvelopeStatus.voided, EnvelopeStatus.declined])).label("failed"),
            func.min(Envelope.created_at).label("submitted"),
        )
        .where(Envelope.user_id == current_user.id)
        .where(Envelope.batch_id.isnot(None))
        .group_by(Envelope.batch_id)
        .order_by(func.min(Envelope.created_at).desc())
    )
    batches = []
    for row in result.all():
        total = row.total
        batches.append({
            "batch_id": row.batch_id,
            "name": row.name,
            "total": total,
            "sent": row.sent,
            "completed": row.completed,
            "failed": row.failed,
            "status": (
                "Partial" if row.failed > 0
                else "Processing" if (row.sent + row.completed + row.failed) < total
                else "Processed"
            ),
            "submitted": row.submitted.isoformat() if row.submitted else None,
        })
    return batches


@router.get("/envelopes/{envelope_id}", response_model=EnvelopeDetailResponse)
async def get_envelope(
    envelope_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    envelope = await svc.get_envelope(db, envelope_id, current_user.id, user_email=current_user.email)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    response = EnvelopeDetailResponse.model_validate(envelope)
    if hasattr(envelope, "owner") and envelope.owner is not None:
        response.from_name = envelope.owner.name
        response.from_email = envelope.owner.email
    else:
        response.from_name = current_user.name
        response.from_email = current_user.email
    return response


@router.put("/envelopes/{envelope_id}", response_model=EnvelopeResponse)
async def update_envelope(
    envelope_id: uuid.UUID,
    data: EnvelopeUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    envelope = await svc.get_envelope(db, envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    if envelope.status != EnvelopeStatus.draft:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Only draft envelopes can be updated",
        )
    return await svc.update_envelope(db, envelope, data)


@router.delete("/envelopes/{envelope_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_envelope(
    envelope_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    envelope = await svc.get_envelope(db, envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    if envelope.status != EnvelopeStatus.draft:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Only draft envelopes can be deleted",
        )
    await svc.delete_envelope(db, envelope)


@router.post("/envelopes/{envelope_id}/send", response_model=EnvelopeDetailResponse)
async def send_envelope(
    envelope_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    envelope = await svc.get_envelope(db, envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    if not envelope.recipients:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Envelope must have at least one recipient before sending",
        )
    if not envelope.documents:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Envelope must have at least one document before sending",
        )
    sent = await svc.send_envelope(db, envelope)
    # Reload with relationships so recipients (including signing_token) are returned
    reloaded = await svc.get_envelope(db, sent.id, current_user.id)
    response = EnvelopeDetailResponse.model_validate(reloaded)
    if hasattr(reloaded, "owner") and reloaded.owner is not None:
        response.from_name = reloaded.owner.name
        response.from_email = reloaded.owner.email
    else:
        response.from_name = current_user.name
        response.from_email = current_user.email
    return response


@router.post("/envelopes/{envelope_id}/void", response_model=EnvelopeResponse)
async def void_envelope(
    envelope_id: uuid.UUID,
    data: VoidEnvelopeRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    envelope = await svc.get_envelope(db, envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    return await svc.void_envelope(db, envelope, data)


@router.post("/envelopes/{envelope_id}/resend", response_model=EnvelopeResponse)
async def resend_envelope(
    envelope_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    envelope = await svc.get_envelope(db, envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    return await svc.resend_envelope(db, envelope)


@router.post(
    "/envelopes/{envelope_id}/save-as-template",
    response_model=SaveAsTemplateResponse,
    status_code=status.HTTP_201_CREATED,
)
async def save_envelope_as_template(
    envelope_id: uuid.UUID,
    data: SaveAsTemplateRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Copy this envelope's documents, fields config, and roles into a new Template record."""
    from app.templates.models import Template

    envelope = await svc.get_envelope(db, envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")

    # Build document_ids list (store as string UUIDs for the template JSON column)
    document_ids = [str(doc.id) for doc in envelope.documents]

    # Build roles list from the envelope's recipients (placeholder roles)
    role_labels = data.role_labels or {}
    roles = []
    for r in envelope.recipients:
        role_entry: dict = {
            "recipient_id": str(r.id),
            "name": r.name,
            "email": r.email,
            "role": r.role.value,
            "routing_order": r.routing_order,
        }
        if r.access_code:
            role_entry["access_code"] = r.access_code
        label = role_labels.get(str(r.id))
        if label:
            role_entry["role_label"] = label
        roles.append(role_entry)

    # Build fields_config from all fields across all documents
    fields_config = [
        {
            "document_id": str(f.document_id),
            "recipient_id": str(f.recipient_id),
            "type": f.type.value,
            "page": f.page,
            "x": f.x,
            "y": f.y,
            "width": f.width,
            "height": f.height,
            "required": f.required,
            "label": f.label,
        }
        for doc in envelope.documents
        for f in doc.fields
    ]

    template_name = data.name or envelope.subject
    template = Template(
        user_id=current_user.id,
        name=template_name,
        description=envelope.message,
        document_ids=document_ids,
        roles=roles,
        fields_config=fields_config,
    )
    db.add(template)
    await db.commit()
    await db.refresh(template)

    return SaveAsTemplateResponse(template_id=template.id)


# ── Document endpoints ────────────────────────────────────────────────────────

@router.post(
    "/envelopes/{envelope_id}/documents",
    response_model=DocumentResponse,
    status_code=status.HTTP_201_CREATED,
)
async def upload_document(
    envelope_id: uuid.UUID,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    envelope = await svc.get_envelope(db, envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    if envelope.status != EnvelopeStatus.draft:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Documents can only be added to draft envelopes",
        )

    # Validate file type — accept PDF and Word documents
    allowed_mime_types = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }
    content_type = file.content_type or ""
    original_filename = file.filename or "document"
    name_lower = original_filename.lower()
    is_pdf = content_type == "application/pdf" or name_lower.endswith(".pdf")
    is_docx = (
        content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or name_lower.endswith(".docx")
    )
    is_doc = content_type == "application/msword" or name_lower.endswith(".doc")

    ext = os.path.splitext(original_filename)[-1].lower()
    if ext not in (".pdf", ".docx", ".doc"):
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported file type: '{ext}'. Accepted types: PDF, DOCX, DOC",
        )
    if content_type and content_type not in allowed_mime_types:
        if not (is_pdf or is_docx or is_doc):
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail=f"Unsupported file type: {content_type}. Accepted types: PDF, DOCX, DOC",
            )

    os.makedirs(settings.upload_dir, exist_ok=True)
    file_id = str(uuid.uuid4())
    stored_filename = f"{file_id}{ext}"
    file_path = os.path.join(settings.upload_dir, stored_filename)

    content = await file.read()

    # Enforce 10MB file size limit
    max_file_size = 10 * 1024 * 1024  # 10MB
    if len(content) > max_file_size:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="File size exceeds 10MB limit",
        )

    with open(file_path, "wb") as f:
        f.write(content)

    file_size = len(content)

    # For .doc/.docx files, convert to a preview PDF for accurate page rendering.
    # The original file is preserved for download; the preview PDF is used for display.
    preview_filename: str | None = None
    if is_pdf:
        try:
            page_count = get_page_count(file_path)
        except Exception:
            page_count = 1
    elif is_docx or is_doc:
        preview_pdf_filename = f"{file_id}_preview.pdf"
        preview_pdf_path = os.path.join(settings.upload_dir, preview_pdf_filename)
        converted = convert_doc_to_pdf(file_path, preview_pdf_path)
        if converted and os.path.exists(preview_pdf_path):
            preview_filename = preview_pdf_filename
            try:
                page_count = get_page_count(preview_pdf_path)
            except Exception:
                page_count = 1
        else:
            # Conversion failed; fall back to section-count heuristic for .docx
            if is_docx:
                try:
                    from docx import Document as DocxDocument
                    docx_doc = DocxDocument(file_path)
                    page_count = max(1, len(docx_doc.sections))
                except Exception:
                    page_count = 1
            else:
                page_count = 1
    else:
        page_count = 1

    order = len(envelope.documents)
    doc = Document(
        envelope_id=envelope_id,
        filename=stored_filename,
        original_filename=original_filename,
        file_path=file_path,
        preview_filename=preview_filename,
        page_count=page_count,
        file_size=file_size,
        order=order,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return {
        "id": doc.id,
        "envelope_id": doc.envelope_id,
        "filename": doc.filename,
        "original_filename": doc.original_filename,
        "page_count": doc.page_count,
        "file_size": doc.file_size,
        "order": doc.order,
        "created_at": doc.created_at,
        "fields": [],
    }


@router.get("/documents/{document_id}/pages/{page}")
async def get_document_page(
    document_id: uuid.UUID,
    page: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = await svc.get_document(db, document_id)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    # Verify ownership via envelope
    envelope = await svc.get_envelope(db, doc.envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

    # Resolve preview PDF path if one was generated during upload
    preview_path: str | None = None
    if doc.preview_filename:
        preview_path = os.path.join(settings.upload_dir, doc.preview_filename)

    try:
        png_bytes = render_page(doc.file_path, page - 1, preview_path=preview_path)
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to render page: {exc}",
        )

    return FastAPIResponse(content=png_bytes, media_type="image/png")


@router.delete("/documents/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    document_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = await svc.get_document(db, document_id)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    envelope = await svc.get_envelope(db, doc.envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")
    if envelope.status != EnvelopeStatus.draft:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Documents can only be removed from draft envelopes",
        )
    # Clean up files from disk
    if doc.file_path and os.path.exists(doc.file_path):
        try:
            os.remove(doc.file_path)
        except OSError:
            pass
    if doc.preview_filename:
        preview_path = os.path.join(settings.upload_dir, doc.preview_filename)
        if os.path.exists(preview_path):
            try:
                os.remove(preview_path)
            except OSError:
                pass
    await svc.delete_document(db, doc)


# ── Recipient endpoints ───────────────────────────────────────────────────────

@router.post(
    "/envelopes/{envelope_id}/recipients",
    response_model=RecipientResponse,
    status_code=status.HTTP_201_CREATED,
)
async def add_recipient(
    envelope_id: uuid.UUID,
    data: RecipientCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    envelope = await svc.get_envelope(db, envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    if envelope.status != EnvelopeStatus.draft:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Recipients can only be added to draft envelopes",
        )
    recipient = await svc.add_recipient(db, envelope, data)
    return {
        "id": recipient.id,
        "envelope_id": recipient.envelope_id,
        "name": recipient.name,
        "email": recipient.email,
        "role": recipient.role,
        "routing_order": recipient.routing_order,
        "status": recipient.status,
        "signing_token": recipient.signing_token,
        "signed_at": recipient.signed_at,
        "declined_at": recipient.declined_at,
        "decline_reason": recipient.decline_reason,
        "access_code": "••••" if recipient.access_code else None,
        "fields": [],
    }


@router.put("/recipients/{recipient_id}", response_model=RecipientResponse)
async def update_recipient(
    recipient_id: uuid.UUID,
    data: RecipientUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    recipient = await svc.get_recipient(db, recipient_id)
    if not recipient:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Recipient not found")
    envelope = await svc.get_envelope(db, recipient.envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")
    if envelope.status != EnvelopeStatus.draft:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Cannot edit recipients on a non-draft envelope")
    updated = await svc.update_recipient(db, recipient, data)
    return {
        "id": updated.id,
        "envelope_id": updated.envelope_id,
        "name": updated.name,
        "email": updated.email,
        "role": updated.role,
        "routing_order": updated.routing_order,
        "status": updated.status,
        "signing_token": updated.signing_token,
        "signed_at": updated.signed_at,
        "declined_at": updated.declined_at,
        "decline_reason": updated.decline_reason,
        "access_code": "••••" if updated.access_code else None,
        "fields": [],
    }


@router.delete("/recipients/{recipient_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_recipient(
    recipient_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    recipient = await svc.get_recipient(db, recipient_id)
    if not recipient:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Recipient not found")
    envelope = await svc.get_envelope(db, recipient.envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")
    if envelope.status != EnvelopeStatus.draft:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Recipients can only be removed from draft envelopes",
        )
    await svc.delete_recipient(db, recipient)


# ── Field endpoints ───────────────────────────────────────────────────────────

@router.post(
    "/documents/{document_id}/fields",
    response_model=FieldResponse,
    status_code=status.HTTP_201_CREATED,
)
async def create_field(
    document_id: uuid.UUID,
    data: FieldCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = await svc.get_document(db, document_id)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    envelope = await svc.get_envelope(db, doc.envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")
    if envelope.status != EnvelopeStatus.draft:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Cannot add fields to a non-draft envelope",
        )
    return await svc.create_field(db, document_id, data)


@router.put("/fields/{field_id}", response_model=FieldResponse)
async def update_field(
    field_id: uuid.UUID,
    data: FieldUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    field = await svc.get_field(db, field_id)
    if not field:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Field not found")
    doc = await svc.get_document(db, field.document_id)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    envelope = await svc.get_envelope(db, doc.envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")
    return await svc.update_field(db, field, data)


@router.delete("/fields/{field_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_field(
    field_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    field = await svc.get_field(db, field_id)
    if not field:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Field not found")
    doc = await svc.get_document(db, field.document_id)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    envelope = await svc.get_envelope(db, doc.envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")
    await svc.delete_field(db, field)


@router.get("/envelopes/{envelope_id}/fields", response_model=list[FieldResponse])
async def get_envelope_fields(
    envelope_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    envelope = await svc.get_envelope(db, envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    return await svc.get_fields_for_envelope(db, envelope_id)


@router.put("/envelopes/{envelope_id}/fields", response_model=list[FieldResponse])
async def save_envelope_fields(
    envelope_id: uuid.UUID,
    data: FieldBulkSaveRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Bulk-save (upsert + delete) all fields for an envelope in one request."""
    envelope = await svc.get_envelope(db, envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    if envelope.status != EnvelopeStatus.draft:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Fields can only be modified on draft envelopes",
        )
    return await svc.save_fields_for_envelope(db, envelope_id, data.fields)


# ── Download endpoints ────────────────────────────────────────────────────────

@router.get("/documents/{document_id}/download")
async def download_document(
    document_id: uuid.UUID,
    view: bool = Query(False),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Download the original document file."""
    doc = await svc.get_document(db, document_id)
    if not doc:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    # Verify ownership via envelope
    envelope = await svc.get_envelope(db, doc.envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

    if not os.path.exists(doc.file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="File not found on server",
        )

    with open(doc.file_path, "rb") as f:
        content = f.read()

    # Determine MIME type from stored filename extension
    _ext = os.path.splitext(doc.filename)[-1].lower()
    _mime_map = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
    }
    media_type = _mime_map.get(_ext, "application/octet-stream")

    disposition = "inline" if view else "attachment"
    safe_name = (doc.original_filename or "document").replace('"', "").replace("\\", "")
    return FastAPIResponse(
        content=content,
        media_type=media_type,
        headers={
            "Content-Disposition": f'{disposition}; filename="{safe_name}"'
        },
    )


@router.get("/envelopes/{envelope_id}/download")
async def download_signed_envelope(
    envelope_id: uuid.UUID,
    view: bool = Query(False),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Download envelope PDF.

    For completed envelopes: returns the signed PDF with all field values applied.
    For non-completed envelopes: returns the original document without field overlays.
    For multi-document envelopes: merges all document PDFs into a single PDF.
    """
    import fitz  # PyMuPDF

    from app.signing.pdf_processor import apply_fields_to_pdf

    envelope = await svc.get_envelope(db, envelope_id, current_user.id, user_email=current_user.email)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")

    if not envelope.documents:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No documents found for this envelope",
        )

    # Verify all document files exist
    for doc in envelope.documents:
        if not os.path.exists(doc.file_path):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Document file not found on server: {doc.original_filename or 'unknown'}",
            )

    safe_subject = "".join(c for c in envelope.subject if c.isalnum() or c in " _-")[:50]

    # Build a PDF for each document (applying fields if completed)
    pdf_parts: list[bytes] = []
    for doc in envelope.documents:
        file_ext = os.path.splitext(doc.file_path)[-1].lower()

        # Resolve the effective PDF path to use for this document.
        # For DOCX/DOC files a preview PDF is generated at upload time; use it
        # when available so we always feed valid PDF bytes to fitz. If no preview
        # exists yet, generate one on-the-fly and fall back to an empty page only
        # as a last resort.
        if file_ext in (".docx", ".doc"):
            if doc.preview_filename:
                preview_pdf_path = os.path.join(settings.upload_dir, doc.preview_filename)
            else:
                preview_pdf_path = ""

            if preview_pdf_path and os.path.exists(preview_pdf_path):
                effective_pdf_path = preview_pdf_path
            else:
                # Generate a preview PDF on-the-fly
                from app.signing.pdf_processor import convert_doc_to_pdf as _cvt
                tmp_preview = os.path.join(
                    settings.upload_dir,
                    f"{os.path.splitext(os.path.basename(doc.file_path))[0]}_preview_dl.pdf",
                )
                converted = _cvt(doc.file_path, tmp_preview)
                effective_pdf_path = tmp_preview if converted and os.path.exists(tmp_preview) else None
        else:
            effective_pdf_path = doc.file_path

        # Collect fields with values for this document
        doc_fields = [
            {
                "type": f.type.value,
                "page": f.page,
                "x": f.x,
                "y": f.y,
                "width": f.width,
                "height": f.height,
                "value": f.value,
                "label": f.label,
            }
            for f in doc.fields
            if f.value
        ]

        if effective_pdf_path and os.path.exists(effective_pdf_path):
            if doc_fields:
                try:
                    pdf_parts.append(apply_fields_to_pdf(effective_pdf_path, doc_fields))
                except Exception:
                    with open(effective_pdf_path, "rb") as fh:
                        pdf_parts.append(fh.read())
            else:
                with open(effective_pdf_path, "rb") as fh:
                    pdf_parts.append(fh.read())
        else:
            # Absolute fallback: generate a minimal blank PDF page via fitz
            fallback_doc = fitz.open()
            fallback_doc.new_page()
            pdf_parts.append(fallback_doc.tobytes())
            fallback_doc.close()

    # If single document, return it directly; otherwise merge with PyMuPDF
    if len(pdf_parts) == 1:
        pdf_bytes = pdf_parts[0]
    else:
        merged = fitz.open()
        for part in pdf_parts:
            src = fitz.open(stream=part, filetype="pdf")
            merged.insert_pdf(src)
            src.close()
        pdf_bytes = merged.tobytes()
        merged.close()

    filename = f"signed_{safe_subject}.pdf" if envelope.status == EnvelopeStatus.completed else f"{safe_subject}.pdf"
    safe_fn = filename.replace('"', "").replace("\\", "")

    disposition = "inline" if view else "attachment"
    return FastAPIResponse(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'{disposition}; filename="{safe_fn}"'},
    )


@router.get("/envelopes/{envelope_id}/certificate")
async def download_certificate(
    envelope_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Download the audit trail / completion certificate PDF."""
    from app.audit import service as audit_svc
    from app.signing.pdf_processor import generate_certificate

    envelope = await svc.get_envelope(db, envelope_id, current_user.id, user_email=current_user.email)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")

    if envelope.status != EnvelopeStatus.completed:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Certificate is only available for completed envelopes",
        )

    audit_events = await audit_svc.get_audit_events(db, envelope_id)

    envelope_data = {
        "id": envelope.id,
        "subject": envelope.subject,
        "sent_at": envelope.sent_at,
        "completed_at": envelope.completed_at,
    }
    audit_dicts = [
        {
            "created_at": ae.created_at,
            "event_type": ae.event_type,
            "details": ae.details,
        }
        for ae in audit_events
    ]

    try:
        cert_bytes = generate_certificate(envelope_data, audit_dicts)
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate certificate: {exc}",
        )

    safe_subject = "".join(c for c in envelope.subject if c.isalnum() or c in " _-")[:50]
    filename = f"certificate_{safe_subject}.pdf"

    return FastAPIResponse(
        content=cert_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Bulk Send endpoints ────────────────────────────────────────────────────────

@router.post(
    "/envelopes/bulk-send-direct",
    response_model=BulkSendResponse,
    status_code=status.HTTP_201_CREATED,
)
async def bulk_send_direct(
    csv_file: UploadFile = File(...),
    documents: List[UploadFile] = File(...),
    subject: str = Form(...),
    message: str = Form(""),
    reminder_days: int = Form(0),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Create one envelope per CSV row (columns: name, email) using the uploaded documents directly."""
    # Parse CSV
    csv_content = await csv_file.read()
    try:
        text = csv_content.decode("utf-8-sig")
        reader = csv.DictReader(io.StringIO(text))
        rows = [row for row in reader]
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Invalid CSV: {exc}",
        )

    if not rows:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="CSV has no data rows",
        )

    # Read all uploaded document bytes once so we can copy them per envelope
    os.makedirs(settings.upload_dir, exist_ok=True)
    source_docs: list[dict] = []
    for upload in documents:
        original_filename = upload.filename or "document"
        ext = os.path.splitext(original_filename)[-1].lower()
        if ext not in (".pdf", ".docx", ".doc"):
            ext = ".pdf"

        file_id = str(uuid.uuid4())
        stored_filename = f"{file_id}{ext}"
        source_path = os.path.join(settings.upload_dir, stored_filename)

        content = await upload.read()
        with open(source_path, "wb") as fh:
            fh.write(content)

        # Determine page count for the source file
        name_lower = original_filename.lower()
        is_pdf = name_lower.endswith(".pdf")
        is_docx = name_lower.endswith(".docx")
        if is_pdf:
            try:
                page_count = get_page_count(source_path)
            except Exception:
                page_count = 1
        elif is_docx:
            try:
                from docx import Document as DocxDocument
                docx_doc = DocxDocument(source_path)
                page_count = max(1, len(docx_doc.sections))
            except Exception:
                page_count = 1
        else:
            page_count = 1

        source_docs.append(
            {
                "source_path": source_path,
                "original_filename": original_filename,
                "ext": ext,
                "file_size": len(content),
                "page_count": page_count,
            }
        )

    batch_id = str(uuid.uuid4())
    created = 0
    failed = 0
    valid_rows = [r for r in rows if (r.get("name") or "").strip() and (r.get("email") or "").strip()]
    total_rows = len(valid_rows)

    for row in valid_rows:
        name = (row.get("name") or "").strip()
        email = (row.get("email") or "").strip()

        # Create envelope
        from app.envelopes.schemas import EnvelopeCreate as EnvCreate
        envelope_data = EnvCreate(
            subject=subject,
            message=message or None,
            reminder_days=reminder_days,
        )
        envelope = await svc.create_envelope(db, current_user.id, envelope_data)
        envelope.batch_id = batch_id
        await db.flush()

        # Copy each document to a new path and create Document records
        for idx, src in enumerate(source_docs):
            new_file_id = str(uuid.uuid4())
            new_stored_filename = f"{new_file_id}{src['ext']}"
            new_file_path = os.path.join(settings.upload_dir, new_stored_filename)
            shutil.copy2(src["source_path"], new_file_path)

            doc = Document(
                envelope_id=envelope.id,
                filename=new_stored_filename,
                original_filename=src["original_filename"],
                file_path=new_file_path,
                page_count=src["page_count"],
                file_size=src["file_size"],
                order=idx,
            )
            db.add(doc)

        await db.flush()

        # Add recipient
        from app.envelopes.models import Recipient as RecipientModel
        recipient = RecipientModel(
            envelope_id=envelope.id,
            name=name,
            email=email,
        )
        db.add(recipient)
        await db.flush()

        # Add default signature field on first document for this recipient
        from app.envelopes.models import Field, FieldType
        first_doc_result = await db.execute(
            select(Document).where(Document.envelope_id == envelope.id).order_by(Document.order).limit(1)
        )
        first_doc = first_doc_result.scalar_one_or_none()
        if first_doc:
            sig_field = Field(
                document_id=first_doc.id,
                recipient_id=recipient.id,
                type=FieldType.signature,
                page=1,
                x=30.0,
                y=85.0,
                width=25.0,
                height=5.0,
                required=True,
            )
            db.add(sig_field)
            await db.flush()

        envelope_detail = await svc.get_envelope(db, envelope.id, current_user.id)
        if envelope_detail:
            try:
                await svc.send_envelope(db, envelope_detail)
                created += 1
            except Exception as exc:
                logger.error("bulk-send-direct: envelope %s failed for %s: %s", envelope.id, email, exc)
                await db.delete(envelope)
                await db.flush()
                failed += 1
        else:
            logger.error("bulk-send-direct: envelope %s not found after create for %s", envelope.id, email)
            failed += 1

    await db.commit()
    return BulkSendResponse(batch_id=batch_id, total=total_rows, created=created, failed=failed)


@router.post(
    "/envelopes/bulk-send",
    response_model=BulkSendResponse,
    status_code=status.HTTP_201_CREATED,
)
async def bulk_send_envelopes(
    template_id: uuid.UUID = Query(...),
    csv_file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Create one envelope per CSV row (columns: name, email) from a template."""
    from app.templates import service as tmpl_svc
    from app.templates.models import Template
    from sqlalchemy import select as sa_select

    # Load template
    tmpl_result = await db.execute(
        sa_select(Template).where(Template.id == template_id, Template.user_id == current_user.id)
    )
    template = tmpl_result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Template not found")

    # Parse CSV
    content = await csv_file.read()
    try:
        text = content.decode("utf-8-sig")
        reader = csv.DictReader(io.StringIO(text))
        rows = [row for row in reader]
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=f"Invalid CSV: {exc}")

    if not rows:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="CSV has no data rows")

    # Validate required columns
    for i, row in enumerate(rows, 1):
        if "name" not in row or "email" not in row:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Row {i} is missing 'name' or 'email' column",
            )

    batch_id = str(uuid.uuid4())
    created = 0
    failed = 0

    for row in rows:
        name = (row.get("name") or "").strip()
        email = (row.get("email") or "").strip()
        if not name or not email:
            continue

        # Create envelope from template
        envelope_id = await tmpl_svc.create_envelope_from_template(db, template, current_user.id)

        # Load the created envelope and set batch_id + add recipient
        env_result = await db.execute(
            sa_select(Envelope).where(Envelope.id == envelope_id)
        )
        envelope = env_result.scalar_one()
        envelope.batch_id = batch_id

        # Delete template-copied recipients to avoid duplicates before adding the CSV recipient
        from app.envelopes.models import Recipient as RecipientModel
        from sqlalchemy import delete as sa_delete
        await db.execute(
            sa_delete(RecipientModel).where(RecipientModel.envelope_id == envelope.id)
        )
        await db.flush()

        recipient = RecipientModel(
            envelope_id=envelope.id,
            name=name,
            email=email,
        )
        db.add(recipient)
        await db.flush()

        from app.envelopes.service import send_envelope
        envelope_detail = await svc.get_envelope(db, envelope.id, current_user.id)
        if envelope_detail:
            try:
                await send_envelope(db, envelope_detail)
                created += 1
            except Exception as exc:
                logger.error("bulk-send: envelope %s failed for %s: %s", envelope.id, email, exc)
                failed += 1
        else:
            logger.error("bulk-send: envelope %s not found after create for %s", envelope.id, email)
            failed += 1

    await db.commit()
    return BulkSendResponse(batch_id=batch_id, total=len(rows), created=created, failed=failed)


# ── Folder endpoints ───────────────────────────────────────────────────────────

@router.get("/folders", response_model=list[FolderResponse])
async def list_folders(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Return the current user's custom folders."""
    folders = await svc.list_folders(db, current_user.id)
    return folders


@router.post("/folders", response_model=FolderResponse, status_code=status.HTTP_201_CREATED)
async def create_folder(
    data: FolderCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Create a new custom folder."""
    if not data.name.strip():
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Folder name cannot be empty",
        )
    folder = await svc.create_folder(db, current_user.id, data.name)
    return folder


@router.delete("/folders/{folder_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_folder(
    folder_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Delete a custom folder. Envelopes in the folder are moved back to no-folder."""
    found = await svc.delete_folder(db, folder_id, current_user.id)
    if not found:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Folder not found")


# ── Comment endpoints ──────────────────────────────────────────────────────────

@router.get("/envelopes/{envelope_id}/comments", response_model=list[CommentResponse])
async def list_comments(
    envelope_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # Allow both the envelope owner and recipients to list comments.
    envelope = await svc.get_envelope(db, envelope_id, current_user.id, user_email=current_user.email)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    if not envelope.allow_comments:
        return []
    comments = await svc.list_comments(db, envelope_id)
    return [
        {
            "id": c.id,
            "envelope_id": c.envelope_id,
            "user_id": c.user_id,
            "text": c.text,
            "created_at": c.created_at,
            "author_name": c.author.name if c.author else None,
            "author_email": c.author.email if c.author else None,
        }
        for c in comments
    ]


@router.post(
    "/envelopes/{envelope_id}/comments",
    response_model=CommentResponse,
    status_code=status.HTTP_201_CREATED,
)
async def create_comment(
    envelope_id: uuid.UUID,
    data: CommentCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # Allow both the envelope owner and recipients to post comments.
    envelope = await svc.get_envelope(db, envelope_id, current_user.id, user_email=current_user.email)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    if not envelope.allow_comments:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Comments are disabled for this envelope",
        )
    comment = await svc.create_comment(db, envelope_id, current_user.id, data)
    return {
        "id": comment.id,
        "envelope_id": comment.envelope_id,
        "user_id": comment.user_id,
        "text": comment.text,
        "created_at": comment.created_at,
        "author_name": comment.author.name if comment.author else None,
        "author_email": comment.author.email if comment.author else None,
    }


@router.delete(
    "/envelopes/{envelope_id}/comments/{comment_id}",
    status_code=status.HTTP_204_NO_CONTENT,
)
async def delete_comment(
    envelope_id: uuid.UUID,
    comment_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # Allow both the envelope owner and recipients to delete their own comments.
    envelope = await svc.get_envelope(db, envelope_id, current_user.id, user_email=current_user.email)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    deleted = await svc.delete_comment(db, comment_id, current_user.id, envelope_id)
    if not deleted:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Comment not found")


# ── Correct in-flight endpoints ────────────────────────────────────────────────

@router.post("/envelopes/{envelope_id}/correct", response_model=EnvelopeDetailResponse)
async def correct_envelope(
    envelope_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Transition a sent/delivered envelope back to draft for editing."""
    envelope = await svc.get_envelope(db, envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    await svc.correct_envelope(db, envelope)
    # Reload with full relationships so the response includes recipients/documents
    reloaded = await svc.get_envelope(db, envelope_id, current_user.id)
    response = EnvelopeDetailResponse.model_validate(reloaded)
    if hasattr(reloaded, "owner") and reloaded.owner is not None:
        response.from_name = reloaded.owner.name
        response.from_email = reloaded.owner.email
    else:
        response.from_name = current_user.name
        response.from_email = current_user.email
    return response


@router.post("/envelopes/{envelope_id}/resend-corrected", response_model=EnvelopeDetailResponse)
async def resend_corrected_envelope(
    envelope_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Re-send an envelope that was corrected (back in draft state)."""
    envelope = await svc.get_envelope(db, envelope_id, current_user.id)
    if not envelope:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Envelope not found")
    await svc.resend_corrected_envelope(db, envelope)
    # Reload with full relationships so the response includes signing tokens
    reloaded = await svc.get_envelope(db, envelope_id, current_user.id)
    response = EnvelopeDetailResponse.model_validate(reloaded)
    if hasattr(reloaded, "owner") and reloaded.owner is not None:
        response.from_name = reloaded.owner.name
        response.from_email = reloaded.owner.email
    else:
        response.from_name = current_user.name
        response.from_email = current_user.email
    return response
